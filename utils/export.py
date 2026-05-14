import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

SECTION_TITLES = {
    "introduction": "1. Introduction",
    "literature_review": "2. Literature Review",
    "methodology": "3. Methodology",
    "results_discussion": "4. Results and Discussion",
    "conclusion": "5. Conclusion",
}


def _clean_text(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip()


def export_docx(title: str, topic: str, sections: dict, references: list, style: str = "apa") -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()
    section_order = ["introduction", "literature_review", "methodology", "results_discussion", "conclusion"]
    for key in section_order:
        content = sections.get(key, "")
        if not content:
            continue
        heading = doc.add_heading(SECTION_TITLES.get(key, key.replace("_", " ").title()), level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        for para_text in [p.strip() for p in content.split("\n\n") if p.strip()]:
            p = doc.add_paragraph(_clean_text(para_text))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.size = Pt(12)
    if references:
        doc.add_heading("References", level=1)
        for ref in references:
            p = doc.add_paragraph(_clean_text(ref))
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.size = Pt(11)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_pdf(title: str, sections: dict, references: list) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=1.25*inch, rightMargin=1.25*inch,
                            topMargin=1*inch, bottomMargin=1*inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=16, spaceAfter=20, alignment=TA_CENTER)
    heading_style = ParagraphStyle("SectionHeading", parent=styles["Heading1"], fontSize=13, spaceBefore=16, spaceAfter=8)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=8, alignment=TA_JUSTIFY)
    ref_style = ParagraphStyle("Ref", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=6, leftIndent=20, firstLineIndent=-20)
    story = [Paragraph(_clean_text(title), title_style), Spacer(1, 12)]
    for key in ["introduction", "literature_review", "methodology", "results_discussion", "conclusion"]:
        content = sections.get(key, "")
        if not content:
            continue
        story.append(Paragraph(SECTION_TITLES.get(key, key.title()), heading_style))
        for para_text in [p.strip() for p in content.split("\n\n") if p.strip()]:
            story.append(Paragraph(_clean_text(para_text), body_style))
    if references:
        story.append(Paragraph("References", heading_style))
        for ref in references:
            story.append(Paragraph(_clean_text(ref), ref_style))
    doc.build(story)
    buffer.seek(0)
    return buffer.read()
